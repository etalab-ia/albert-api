from typing import List, Optional

from docx import Document as DocxLoader
from langchain.docstore.document import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter


class DocxParser:
    def __init__(self):
        pass

    def parse(
        self, file: DocxLoader, file_name: str, chunk_size: int, chunk_overlap: int, chunk_min_size: Optional[int] = None
    ) -> List[LangchainDocument]:
        """
        Parse a DOCX file and converts it into a list of text chunks.

        Args:
            file (DocxLoader): DOCX file to be processed.
            file_name (str): File name.
            chunk_size (int): Maximum size of each text chunk.
            chunk_overlap (int): Number of characters overlapping between chunks.
            chunk_min_size (int): Minimum size of a chunk to be considered valid.

        Returns:
            List[LangchainDocument]: List of Langchain documents, where each document corresponds to a text chunk.
        """
        documents = []
        chunker = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, length_function=len, is_separator_regex=False)

        title = None
        text_chunks = []

        for paragraph in file.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                if title:
                    full_text = "\n".join([p.text for p in text_chunks])
                    chunks = chunker.split_text(full_text)
                    for k, text in enumerate(chunks):
                        if chunk_min_size and len(text) < chunk_min_size:
                            continue
                        document = LangchainDocument(page_content=self.cleaner.clean_string(text), metadata={"file_id": file_name, "title": title})
                        documents.append(document)
                # Updating title for new subpart
                title = paragraph.text.strip()
                text_chunks = []
            else:
                text_chunks.append(paragraph)

        # Adding the last subpart
        if title:
            full_text = "\n".join([p.text for p in text_chunks])
            splitted_text = chunker.split_text(full_text)

            if splitted_text:
                for k, text in enumerate(splitted_text):
                    if chunk_min_size:
                        if len(text) > chunk_min_size:  # We avoid meaningless little chunks
                            chunk = LangchainDocument(page_content=self.cleaner.clean_string(text), metadata={"file_id": file_name, "title": title})
                    else:
                        chunk = LangchainDocument(page_content=self.cleaner.clean_string(text), metadata={"file_id": file_name, "title": title})
                    documents.append(chunk)

        elif text_chunks:
            full_text = "\n".join([p.text for p in text_chunks])
            chunks = chunker.split_text(full_text)

            for chunk in enumerate(chunks):
                if chunk_min_size and len(chunk) > chunk_min_size:
                    continue

                document = LangchainDocument(page_content=self.cleaner.clean_string(text), metadata={"file_id": file_name, "title": title})
                documents.append(document)

        return documents
